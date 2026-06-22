"""
Generates the RetailAI written application (IBM Generative AI badge / credential)
as a .docx file, structured to match the submission guideline sections.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        if color:
            run.font.color.rgb = RGBColor(*color)
        run.font.bold = True
    return heading

def add_section_title(doc, text):
    """Styled section title with a blue accent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)   # IBM-ish blue
    return p

def add_body(doc, text, bold=False, space_after=Pt(6)):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = space_after
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

def set_page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ──────────────────────────────────────────────────────────────────────────────
# Build document
# ──────────────────────────────────────────────────────────────────────────────

doc = Document()
set_page_margins(doc)

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("Written Application")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(16)
r2 = sub.add_run("IBM Consulting — Generative AI Practitioner Credential")
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)

add_horizontal_rule(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 1 — Generative AI Client or Project
# ──────────────────────────────────────────────────────────────────────────────
add_section_title(doc, "1. Generative AI Client or Project")
add_horizontal_rule(doc)

add_body(doc, (
    "Project Name: RetailAI — Multi-Store Retail Intelligence Platform\n"
    "Project Type: Internal / Demo Application\n"
    "Project Period: 2025 (Development & Demo Deployment)"
), space_after=Pt(8))

add_body(doc, (
    "RetailAI is a fully self-contained, AI-powered retail management platform built to address "
    "a real-world business challenge: how can small and medium retail store owners gain intelligent, "
    "actionable insights from their own sales and inventory data — without relying on expensive cloud "
    "platforms or third-party AI services?"
))

add_body(doc, (
    "The platform was designed to simulate a realistic multi-store retail environment spanning three "
    "distinct business verticals — general retail, consumer electronics, and automotive dealership — "
    "with a shared owner having access to multiple stores from a single login. The business problem "
    "addressed includes:"
))

for point in [
    "Stockout losses due to lack of demand foresight (e.g., beverages selling out in summer, electronics during festivals)",
    "Overstock and perishable wastage caused by purchasing without data-driven guidance",
    "Missed revenue opportunities around seasonal spikes, festivals, and weather events",
    "Lack of store-specific, contextually filtered AI recommendations for retail owners who do not have data science teams",
]:
    add_bullet(doc, point)

add_body(doc, (
    "RetailAI was developed as a complete working demonstration delivering end-to-end generative "
    "AI-adjacent outcomes: from data ingestion and pattern recognition through to explainable, "
    "natural-language-style recommendations surfaced directly in the owner's dashboard."
))

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 2 — watsonx or Strategic Partner Experience
# ──────────────────────────────────────────────────────────────────────────────
add_section_title(doc, "2. watsonx or Strategic Partner Experience")
add_horizontal_rule(doc)

add_body(doc, (
    "RetailAI was built using IBM Bob — IBM Consulting's internal AI-powered software engineering "
    "assistant — as the primary generative AI capability throughout the development lifecycle."
))

add_body(doc, "IBM Bob was used to achieve generative AI outcomes across the following dimensions:")

for point in [
    "Prompt Engineering & Code Generation: The entire application — backend (Python/Flask), frontend (React), AI engine, database schema, seed data, and documentation — was generated from a single structured natural-language prompt authored in IBM Bob. The prompt (saved as Used_prompt.md within the project) represents a comprehensive specification document that drove 100% of the initial code output.",
    "Iterative Refinement: Subsequent enhancements — including financial year-aware GST calculations, multi-store AI insights filtering, Excel upload and template generation, password policy enforcement, and UI/UX polish — were all implemented through follow-up conversational prompts within IBM Bob's Agent mode.",
    "Documentation Generation: All project documentation, including README.md and IMPLEMENTATION_GUIDE.md, was generated and maintained by IBM Bob based on prompts describing documentation requirements.",
    "Debugging & Problem Solving: IBM Bob was used to diagnose and resolve issues including CORS configuration, SQLite query design, React state management, and Chart.js rendering edge cases.",
    "IBM Consulting Alignment: The use of IBM Bob directly embodies IBM Consulting's generative AI-first delivery model, demonstrating how a practitioner can accelerate full-stack application delivery from days/weeks to hours using structured prompting and AI-generated code.",
]:
    add_bullet(doc, point)

add_body(doc, (
    "No AWS, Azure, or other Strategic Partner services were required. The entire generative AI "
    "workflow was executed through IBM Bob, keeping the application fully offline and the development "
    "process fully within IBM's internal AI tooling ecosystem."
))

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 3 — Application, Work Product, or Deliverable
# ──────────────────────────────────────────────────────────────────────────────
add_section_title(doc, "3. Application, Work Product, or Deliverable")
add_horizontal_rule(doc)

add_body(doc, (
    "RetailAI is a complete, working web application delivered as a demo-ready codebase. "
    "The following work products were created:"
))

add_body(doc, "3.1  Backend — Python / Flask API Server", bold=True)
for point in [
    "app.py: Full REST API with 15+ endpoints covering authentication, store management, inventory, dashboard analytics, AI insights, and Excel upload/download.",
    "models.py: SQLite schema with 12 relational tables — Owners, Stores, Products, Customers, Orders, OrderItems, Payments, Inventory, Suppliers, WeatherHistory, FestivalCalendar, DemandSignals.",
    "seed_data.py: Automated data generator producing 3+ years (Jan 2023 – Mar 2026) of realistic retail sales data — over 41,000 orders across 3 stores.",
    "modules/ai_engine.py: Rule-based + statistical AI engine with 7 insight categories including stockout loss estimation, seasonal demand spikes, festival opportunity alerts, weather impact analysis, and overstock warnings.",
    "modules/auth.py: SHA-256 password hashing, UUID token management, and password policy enforcement.",
]:
    add_bullet(doc, point)

add_body(doc, "3.2  Frontend — React Single Page Application", bold=True)
for point in [
    "Dashboard: Financial Year-aware revenue and GST summary (current FY YTD + last 2 completed FYs), compact Indian notation (₹5K / ₹31L / ₹3.2Cr), Chart.js bar and doughnut charts.",
    "Inventory Page: Real-time stock table with inline quantity adjustment, reorder alerts, and overstock badges.",
    "AI Insights Page: Combined recommendations across all stores owned by the logged-in user, with store-level and type-level filtering. Insight cards display priority, category, impact estimate, and actionable recommendation text.",
    "Settings Page: Add Store, Add Product (persisted to DB), and Excel upload with drag-and-drop. Includes a downloadable RetailAI Template.xlsx with Products, Sales History, and Instructions sheets.",
    "Auth Pages: Login and Register with real-time password policy validation (8 chars, uppercase, lowercase, digit, special character) enforced on both frontend and backend.",
    "Multi-store switching: Store selector in sidebar — switching updates all pages dynamically for the selected store.",
]:
    add_bullet(doc, point)

add_body(doc, "3.3  Pre-built Demo Data Asset", bold=True)
add_body(doc, (
    "A pre-built Excel file (RetailAI template Sports.xlsx) containing 38 sports products across "
    "12 categories and over 7,400 generated sales rows (Apr 2024 – present) is provided as a demo "
    "data asset. A new user can register a Sports retail store, upload this file via Settings, and "
    "immediately unlock AI insights — demonstrating the full end-to-end generative AI workflow in under "
    "5 minutes during a live demo."
))

add_body(doc, "3.4  Technical Specifications", bold=True)
specs = [
    ("Backend Language", "Python 3.14.x (minimum 3.9+)"),
    ("Web Framework", "Flask 3.1.3 + Flask-CORS 6.0.5"),
    ("Database", "SQLite (embedded, no external server)"),
    ("Excel Processing", "openpyxl 3.1.5"),
    ("Frontend Framework", "React 19.2.7"),
    ("Charts", "Chart.js 4.5.1 / react-chartjs-2 5.3.1"),
    ("HTTP Client", "Axios 1.18.0"),
    ("Runtime", "Node.js 24.x"),
    ("AI Engine", "Custom rule-based + statistical (no ML libraries, no external APIs)"),
    ("Deployment", "Fully offline — localhost only"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Component"
hdr[1].text = "Technology / Version"
for hcell in hdr:
    for para in hcell.paragraphs:
        for run in para.runs:
            run.bold = True
for label, value in specs:
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value
doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 4 — Approach, Methods, and Tools
# ──────────────────────────────────────────────────────────────────────────────
add_section_title(doc, "4. Approach, Methods, and Tools")
add_horizontal_rule(doc)

add_body(doc, (
    "The development of RetailAI followed an AI-first delivery approach anchored in IBM Consulting's "
    "generative AI methodology, using IBM Bob as the primary generative AI tool."
))

add_body(doc, "4.1  Structured Prompt Engineering (IBM Bob)", bold=True)
add_body(doc, (
    "A single, comprehensive prompt document was authored as the system specification. This prompt — "
    "saved as Used_prompt.md within the project — described the full application in natural language: "
    "business objectives, technical stack, database schema, AI engine behaviour, UI requirements, "
    "authentication rules, and edge cases. IBM Bob's Agent mode processed this prompt and generated "
    "the complete initial codebase in a single session, demonstrating the power of prompt-driven "
    "software delivery."
))

add_body(doc, "4.2  Iterative Delivery via Conversational AI", bold=True)
add_body(doc, (
    "After the initial code generation, subsequent features and refinements were delivered through "
    "follow-up conversational prompts in IBM Bob — mirroring an agile sprint cycle but compressed "
    "into hours. Each iteration included: describe the requirement → AI generates the code → "
    "review and test → prompt for correction if needed."
))

add_body(doc, "4.3  IBM Core Method Alignment", bold=True)
for point in [
    "Design Thinking: User stories were embedded in the prompt (e.g., 'Mr. A. Sharma logs in once and switches between ABC Store and MNO Motors') to ensure the AI generated user-centric solutions.",
    "Agile Delivery: The application was built incrementally — core auth and data model first, followed by AI engine, then UI, then advanced features like Excel upload and GST reporting.",
    "AI Garage Principles: The project demonstrates a minimum viable AI product (MVP) pattern — a working AI-powered solution scoped to a specific business problem, deliverable on a single machine without cloud dependencies.",
]:
    add_bullet(doc, point)

add_body(doc, "4.4  Generative AI Outcome Validation", bold=True)
add_body(doc, (
    "The AI engine was validated against 7 defined demo scenarios (e.g., summer beverage stockout, "
    "festival electronics shortage, rainy weather overstock risk, Q4 vehicle demand surge). Each "
    "scenario was seeded into the database and verified to produce the correct AI recommendation "
    "in the Insights panel — confirming that the generative AI tooling produced a functionally "
    "correct and business-relevant outcome."
))

# ──────────────────────────────────────────────────────────────────────────────
# Footer
# ──────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
add_horizontal_rule(doc)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(6)
fr = footer_p.add_run("Made with IBM Bob")
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)

# ──────────────────────────────────────────────────────────────────────────────
# Save
# ──────────────────────────────────────────────────────────────────────────────
output_path = "RetailAI_Written_Application.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
